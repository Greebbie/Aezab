"""Knowledge management API — upload, manage, search."""

from __future__ import annotations

import logging
import re
from typing import Any

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from server.config import settings
from server.db import get_db
from server.models.knowledge import KnowledgeSource, KnowledgeChunk
from server.schemas.knowledge import (
    KnowledgeSourceCreate, KnowledgeSourceOut,
    KVEntityCreate, FAQCreate,
    RetrievalRequest, RetrievalResponse,
)
from server.engine.knowledge_retriever import KnowledgeRetriever
from server.middleware.auth import get_current_user, get_tenant_id
from server.api._usage_check import get_resource_usage

logger = logging.getLogger(__name__)

router = APIRouter(dependencies=[Depends(get_current_user)])


def _embed_chunk_background(chunk_id: str, text: str, domain: str) -> None:
    """Best-effort vector indexing that must not block API responses."""
    try:
        from server.engine.vector_store import get_vector_store
        vs = get_vector_store()
        if vs:
            vs.add(chunk_id, text, domain=domain)
            vs.save()
    except Exception as exc:
        logger.warning("Background vector indexing failed for chunk %s: %s", chunk_id, exc)


def _embed_batch_background(items: list[dict[str, str]]) -> None:
    """Best-effort batch vector indexing that runs after upload responds."""
    if not items:
        return
    try:
        from server.engine.vector_store import get_vector_store
        vs = get_vector_store()
        if vs:
            vs.add_batch(items)
            vs.save()
    except Exception as exc:
        logger.warning("Background vector batch indexing failed for %d chunks: %s", len(items), exc)


def _delete_vectors_background(chunk_ids: list[str], source_id: str) -> None:
    """Best-effort vector deletion that must not block source deletion."""
    if not chunk_ids:
        return
    try:
        from server.engine.vector_store import get_vector_store
        vs = get_vector_store()
        if vs:
            removed = vs.delete(chunk_ids)
            if removed:
                vs.save()
    except Exception as exc:
        logger.warning(
            "Deleted knowledge source %s but failed to prune vector index: %s",
            source_id,
            exc,
        )


# ── Knowledge Sources ────────────────────────────────────────────

@router.get("/sources", response_model=list[KnowledgeSourceOut])
async def list_sources(
    tenant_id: str = Depends(get_tenant_id),
    domain: str | None = None,
    db: AsyncSession = Depends(get_db),
):
    stmt = select(KnowledgeSource).where(KnowledgeSource.tenant_id == tenant_id)
    if domain:
        stmt = stmt.where(KnowledgeSource.domain == domain)
    result = await db.execute(stmt)
    return result.scalars().all()


@router.post("/sources", response_model=KnowledgeSourceOut, status_code=201)
async def create_source(
    body: KnowledgeSourceCreate,
    tenant_id: str = Depends(get_tenant_id),
    db: AsyncSession = Depends(get_db),
):
    source = KnowledgeSource(
        name=body.name,
        source_type=body.source_type,
        domain=body.domain,
        tenant_id=tenant_id,
        metadata_=body.metadata,
    )
    db.add(source)
    await db.commit()
    await db.refresh(source)
    return source


async def _get_owned_source(
    db: AsyncSession, source_id: str, tenant_id: str,
) -> KnowledgeSource:
    """Fetch a KnowledgeSource, 404ing if missing or owned by another tenant.

    KnowledgeChunk rows have no tenant_id column of their own — ownership is
    always derived by joining through their parent KnowledgeSource.
    """
    result = await db.execute(
        select(KnowledgeSource).where(
            KnowledgeSource.id == source_id, KnowledgeSource.tenant_id == tenant_id
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(404, "Source not found")
    return source


@router.get("/sources/{source_id}/chunks")
async def list_chunks(
    source_id: str, tenant_id: str = Depends(get_tenant_id), db: AsyncSession = Depends(get_db),
):
    await _get_owned_source(db, source_id, tenant_id)
    result = await db.execute(
        select(KnowledgeChunk).where(KnowledgeChunk.source_id == source_id)
        .order_by(KnowledgeChunk.chunk_index)
    )
    chunks = result.scalars().all()
    return [
        {
            "id": c.id,
            "entity_key": c.entity_key,
            "content": c.content,
            "domain": c.domain,
            "chunk_index": c.chunk_index,
        }
        for c in chunks
    ]


@router.get("/sources/{source_id}/usage")
async def get_source_usage(
    source_id: str, tenant_id: str = Depends(get_tenant_id), db: AsyncSession = Depends(get_db),
):
    """Report which agents currently depend on this knowledge source."""
    await _get_owned_source(db, source_id, tenant_id)
    return await get_resource_usage(
        db, tenant_id,
        skill_type="knowledge_qa",
        matches=lambda ec: source_id in (ec.get("knowledge_source_ids") or []),
    )


@router.delete("/sources/{source_id}", status_code=204)
async def delete_source(
    source_id: str,
    background_tasks: BackgroundTasks,
    tenant_id: str = Depends(get_tenant_id),
    db: AsyncSession = Depends(get_db),
):
    source = await _get_owned_source(db, source_id, tenant_id)

    chunks = await db.execute(select(KnowledgeChunk).where(KnowledgeChunk.source_id == source_id))
    chunk_rows = chunks.scalars().all()
    chunk_ids = [chunk.id for chunk in chunk_rows if chunk.id]

    for chunk in chunk_rows:
        await db.delete(chunk)
    await db.delete(source)
    await db.commit()

    background_tasks.add_task(_delete_vectors_background, chunk_ids, source_id)


# ── KV Entities (fast-answer channel) ────────────────────────────

@router.post("/kv", status_code=201)
async def add_kv_entity(
    body: KVEntityCreate,
    background_tasks: BackgroundTasks,
    tenant_id: str = Depends(get_tenant_id),
    db: AsyncSession = Depends(get_db),
):
    source = await _get_owned_source(db, body.source_id, tenant_id)

    chunk = KnowledgeChunk(
        source_id=body.source_id,
        content=body.content,
        entity_key=body.entity_key,
        domain=body.domain,
        metadata_=body.metadata,
    )
    db.add(chunk)
    # Update source chunk count
    source.chunk_count = (source.chunk_count or 0) + 1
    await db.flush()
    await db.commit()

    background_tasks.add_task(_embed_chunk_background, chunk.id, body.content, body.domain)

    return {"id": chunk.id, "entity_key": body.entity_key}


# ── FAQ entries ──────────────────────────────────────────────────

@router.post("/faq", status_code=201)
async def add_faq(
    body: FAQCreate,
    background_tasks: BackgroundTasks,
    tenant_id: str = Depends(get_tenant_id),
    db: AsyncSession = Depends(get_db),
):
    source = await _get_owned_source(db, body.source_id, tenant_id)

    chunk = KnowledgeChunk(
        source_id=body.source_id,
        content=body.answer,
        entity_key=body.question,
        domain=body.domain,
        metadata_={"question": body.question, **(body.metadata or {})},
    )
    db.add(chunk)
    source.chunk_count = (source.chunk_count or 0) + 1
    await db.flush()
    await db.commit()

    background_tasks.add_task(_embed_chunk_background, chunk.id, body.answer, body.domain)

    return {"id": chunk.id, "question": body.question}


# ── Retrieval test ───────────────────────────────────────────────

@router.post("/search", response_model=RetrievalResponse)
async def search(body: RetrievalRequest, db: AsyncSession = Depends(get_db)):
    """Test knowledge retrieval — used from the console for debugging."""
    from server.runtime_config import runtime_config
    vector_store = None
    if body.use_rag_channel:
        try:
            from server.engine.vector_store import get_vector_store_if_initialized
            vector_store = get_vector_store_if_initialized()
        except Exception as exc:
            logger.warning("Vector store unavailable for search: %s", exc)
    retriever = KnowledgeRetriever(
        db, vector_store=vector_store,
        runtime_cfg=runtime_config.all(),
    )
    return await retriever.retrieve(
        query=body.query,
        domain=body.domain,
        top_k=body.top_k,
        use_fast=body.use_fast_channel,
        use_rag=body.use_rag_channel,
    )


# ── Document Upload ─────────────────────────────────────────────

ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx"}
_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030")
_UTF16_ENCODINGS = ("utf-16", "utf-16-le", "utf-16-be")
_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _max_upload_bytes() -> int:
    return max(settings.knowledge_max_upload_mb, 1) * 1024 * 1024


def _sanitize_filename(filename: str) -> str:
    """Strip directory components and reject path-traversal attempts."""
    import os
    # Take only the basename to block directory traversal
    name = os.path.basename(filename)
    # Reject anything that still looks suspicious
    if ".." in name or "/" in name or "\\" in name:
        raise HTTPException(400, "Invalid filename")
    return name


async def _unique_source_name(db: AsyncSession, base_name: str, tenant_id: str) -> str:
    """Return a tenant-unique source name derived from the uploaded filename."""
    clean_name = (base_name or "Uploaded document").strip()[:220]
    result = await db.execute(
        select(KnowledgeSource.name).where(KnowledgeSource.tenant_id == tenant_id)
    )
    existing_names = set(result.scalars().all())
    if clean_name not in existing_names:
        return clean_name

    for suffix in range(2, 1000):
        candidate = f"{clean_name} ({suffix})"
        if candidate not in existing_names:
            return candidate

    raise HTTPException(409, "Unable to generate a unique knowledge source name")


def _decode_text_bytes(raw_bytes: bytes, file_label: str) -> str:
    """Decode user-provided text while rejecting likely binary content."""
    encodings = list(_TEXT_ENCODINGS)
    if _has_utf16_bom(raw_bytes):
        encodings = list(_UTF16_ENCODINGS) + encodings
    elif _looks_like_utf16(raw_bytes):
        encodings += list(_UTF16_ENCODINGS)

    for encoding in encodings:
        try:
            text = raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
        if _looks_like_binary_text(text):
            continue
        cleaned = _clean_extracted_text(text)
        return _require_extracted_text(
            cleaned,
            f"{file_label} is empty after text extraction",
        )

    raise HTTPException(
        400,
        f"{file_label} must be plain text encoded as UTF-8, GB18030, or UTF-16",
    )


def _has_utf16_bom(raw_bytes: bytes) -> bool:
    return raw_bytes.startswith((b"\xff\xfe", b"\xfe\xff"))


def _looks_like_utf16(raw_bytes: bytes) -> bool:
    sample = raw_bytes[:2048]
    if len(sample) < 4:
        return False
    even_positions = sample[0::2]
    odd_positions = sample[1::2]
    even_nul_ratio = even_positions.count(0) / max(len(even_positions), 1)
    odd_nul_ratio = odd_positions.count(0) / max(len(odd_positions), 1)
    return max(even_nul_ratio, odd_nul_ratio) > 0.25


def _clean_extracted_text(text: str) -> str:
    text = text.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n")
    text = _CONTROL_CHAR_RE.sub(" ", text)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def _looks_like_binary_text(text: str) -> bool:
    if not text:
        return False
    control_count = len(_CONTROL_CHAR_RE.findall(text))
    return control_count / max(len(text), 1) > 0.02


def _require_extracted_text(text: str, empty_message: str) -> str:
    cleaned = _clean_extracted_text(text)
    if not cleaned.strip():
        raise HTTPException(400, empty_message)
    return cleaned


def _extract_text_from_pdf(raw_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(
            400,
            "PDF support requires pypdf. Install with: pip install 'aezab[rag]'",
        )
    import io

    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise HTTPException(400, "Invalid PDF file") from exc
    if reader.is_encrypted:
        try:
            decrypted = reader.decrypt("")
        except Exception as exc:
            raise HTTPException(400, "Encrypted PDF files are not supported") from exc
        if not decrypted:
            raise HTTPException(400, "Encrypted PDF files are not supported")

    pages = []
    try:
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    except Exception as exc:
        raise HTTPException(400, "Unable to extract text from PDF file") from exc
    return _require_extracted_text(
        "\n\n".join(pages),
        "PDF file contains no extractable text. Run OCR before uploading scanned PDFs.",
    )


def _extract_text_from_docx(raw_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise HTTPException(
            400,
            "DOCX support requires python-docx. Install with: pip install 'aezab[rag]'",
        )
    import io

    try:
        doc = DocxDocument(io.BytesIO(raw_bytes))
    except Exception as exc:
        raise HTTPException(400, "Invalid DOCX file") from exc

    sections = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for index, table in enumerate(doc.tables, start=1):
        parsed_rows: list[tuple[int, list[str]]] = []
        for row_number, row in enumerate(table.rows, start=1):
            cells = _normalize_table_row(cell.text for cell in row.cells)
            if cells:
                parsed_rows.append((row_number, cells))
        rows = _format_table_rows(parsed_rows)
        if rows:
            sections.append(f"[Table {index}]\n" + "\n".join(rows))
    return _require_extracted_text(
        "\n\n".join(sections),
        "DOCX file contains no extractable text. Use real text instead of screenshots.",
    )


def _extract_text_from_csv(raw_bytes: bytes) -> str:
    """Extract readable text from CSV bytes."""
    import csv
    import io

    text = _decode_text_bytes(raw_bytes, "CSV file")

    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel

    parsed_rows: list[tuple[int, list[str]]] = []
    reader = csv.reader(io.StringIO(text), dialect)
    for row_number, row in enumerate(reader, start=1):
        cells = _normalize_table_row(row)
        if not cells:
            continue
        parsed_rows.append((row_number, cells))
    return _require_extracted_text(
        "\n".join(_format_table_rows(parsed_rows)),
        "CSV file contains no readable rows",
    )


def _extract_text_from_xlsx(raw_bytes: bytes) -> str:
    """Extract readable text from XLSX bytes using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise HTTPException(
            400,
            "XLSX support requires openpyxl. Install with: pip install 'aezab[rag]'",
        )

    import io

    try:
        workbook = load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    except Exception as exc:
        raise HTTPException(400, "Invalid XLSX file") from exc

    try:
        sections: list[str] = []
        for sheet in workbook.worksheets:
            parsed_rows: list[tuple[int, list[str]]] = []
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = _normalize_table_row(row)
                if not cells:
                    continue
                parsed_rows.append((row_number, cells))
            rows = _format_table_rows(parsed_rows)
            if rows:
                sections.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
        return _require_extracted_text(
            "\n\n".join(sections),
            "XLSX file contains no readable cells",
        )
    finally:
        workbook.close()


def _extract_text_by_extension(raw_bytes: bytes, ext: str) -> str:
    if ext == ".pdf":
        return _extract_text_from_pdf(raw_bytes)
    if ext == ".docx":
        return _extract_text_from_docx(raw_bytes)
    if ext == ".csv":
        return _extract_text_from_csv(raw_bytes)
    if ext == ".xlsx":
        return _extract_text_from_xlsx(raw_bytes)
    if ext in {".txt", ".md"}:
        return _decode_text_bytes(raw_bytes, "Text file")
    raise HTTPException(
        400,
        f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
    )


def _normalize_table_row(row: Any) -> list[str]:
    """Convert a CSV/XLSX row to stripped cell text while preserving columns."""
    cells: list[str] = []
    for value in row:
        if value is None:
            cells.append("")
            continue
        text = str(value).strip()
        cells.append(text)
    while cells and not cells[-1]:
        cells.pop()
    return cells


def _format_table_row(cells: list[str], headers: list[str] | None, row_number: int) -> str:
    """Format a table row as header-aware text for retrieval."""
    if not any(cells):
        return ""
    if not headers:
        pairs = [
            f"Column {idx + 1}: {cell}"
            for idx, cell in enumerate(cells)
            if cell
        ]
        return f"Row {row_number}: " + " | ".join(pairs)

    pairs: list[str] = []
    for idx, cell in enumerate(cells):
        if not cell:
            continue
        header = (
            headers[idx].strip()
            if idx < len(headers) and headers[idx].strip()
            else f"Column {idx + 1}"
        )
        pairs.append(f"{header}: {cell}")
    return f"Row {row_number}: " + " | ".join(pairs) if pairs else ""


def _format_table_rows(parsed_rows: list[tuple[int, list[str]]]) -> list[str]:
    if not parsed_rows:
        return []

    _, first_cells = parsed_rows[0]
    headers = first_cells if _looks_like_header_row(first_cells) else None
    rows: list[str] = []
    start_index = 0

    if headers:
        rows.append("Headers: " + " | ".join(headers))
        start_index = 1

    for row_number, cells in parsed_rows[start_index:]:
        line = _format_table_row(
            cells,
            headers=headers,
            row_number=row_number,
        )
        if line:
            rows.append(line)
    return rows


def _looks_like_header_row(cells: list[str]) -> bool:
    non_empty = [cell.strip() for cell in cells if cell.strip()]
    if len(non_empty) < 2:
        return False
    if len({cell.casefold() for cell in non_empty}) != len(non_empty):
        return False
    if any(len(cell) > 40 for cell in non_empty):
        return False
    return not any(_looks_like_table_value(cell) for cell in non_empty)


def _looks_like_table_value(cell: str) -> bool:
    value_markers = "-/\\:@$%0123456789"
    return any(marker in cell for marker in value_markers)


# Sentence-ending punctuation for recursive splitting
_SENTENCE_DELIMITERS = "。.！!？?；;"


def _find_sentence_boundary(text: str, target: int, window: int = 50) -> int:
    """Find the nearest sentence boundary around `target` position.

    Looks within [target - window, target + window] for sentence delimiters.
    Falls back to target if no boundary found.
    """
    start = max(0, target - window)
    end = min(len(text), target + window)
    region = text[start:end]

    # Search for the closest delimiter to the midpoint of the region
    mid = target - start
    best = -1
    best_dist = window + 1

    for i, ch in enumerate(region):
        if ch in _SENTENCE_DELIMITERS:
            dist = abs(i - mid)
            if dist < best_dist:
                best_dist = dist
                best = start + i + 1  # split after the delimiter

    # Also try jieba word boundaries
    if best == -1:
        try:
            import jieba
            tokens = jieba.lcut(region)
            pos = start
            for tok in tokens:
                pos += len(tok)
                dist = abs(pos - target)
                if dist < best_dist:
                    best_dist = dist
                    best = pos
        except ImportError:
            pass

    return best if best != -1 else target


def _recursive_split(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Recursively split text using a hierarchy of separators.

    Separator hierarchy: paragraph (\\n\\n) -> line (\\n) -> sentence -> character.
    Overlap cuts prefer sentence boundaries via _find_sentence_boundary.
    """
    # Level 1: Split by paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    chunks: list[str] = []
    current_chunk = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            # Flush accumulated content
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
            # Level 2: Split long paragraph by lines
            lines = [ln.strip() for ln in paragraph.split("\n") if ln.strip()]
            line_chunk = ""
            for line in lines:
                if len(line) > chunk_size:
                    if line_chunk:
                        chunks.append(line_chunk)
                        line_chunk = ""
                    # Level 3: Split long line at sentence boundaries
                    start = 0
                    while start < len(line):
                        end = start + chunk_size
                        if end < len(line):
                            boundary = _find_sentence_boundary(line, end)
                            end = boundary
                        chunks.append(line[start:end].strip())
                        # Overlap: step back but align to sentence boundary
                        next_start = end - chunk_overlap
                        if chunk_overlap > 0 and next_start > start:
                            next_start = _find_sentence_boundary(
                                line, next_start, window=chunk_overlap // 2 or 20,
                            )
                        start = max(next_start, start + 1)  # ensure progress
                    continue

                candidate = (line_chunk + "\n" + line).strip() if line_chunk else line
                if len(candidate) > chunk_size and line_chunk:
                    chunks.append(line_chunk)
                    line_chunk = line
                else:
                    line_chunk = candidate

            if line_chunk:
                # Check if it fits with current_chunk
                chunks.append(line_chunk)
            continue

        # Normal-sized paragraph: accumulate
        candidate = (current_chunk + "\n\n" + paragraph).strip() if current_chunk else paragraph
        if len(candidate) > chunk_size and current_chunk:
            chunks.append(current_chunk)
            # Overlap from previous chunk at sentence boundary
            if chunk_overlap > 0 and len(current_chunk) > chunk_overlap:
                overlap_start = len(current_chunk) - chunk_overlap
                overlap_start = _find_sentence_boundary(
                    current_chunk, overlap_start, window=chunk_overlap // 2 or 20,
                )
                current_chunk = current_chunk[overlap_start:] + "\n\n" + paragraph
            else:
                current_chunk = paragraph
        else:
            current_chunk = candidate

    if current_chunk:
        chunks.append(current_chunk)

    return [c for c in chunks if c.strip()]


def _extract_entity_key(content: str) -> str:
    """Auto-generate entity_key from the first sentence of a chunk."""
    # Try splitting by common sentence terminators
    for delimiter in ["。", ".", "！", "!", "？", "?", "\n"]:
        idx = content.find(delimiter)
        if 0 < idx < 200:
            return content[: idx + 1].strip()
    # Fallback: first 100 characters
    return content[:100].strip()


@router.post("/upload", status_code=201)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    source_id: str | None = Form(None),
    source_name: str | None = Form(None),
    domain: str = Form("default"),
    chunk_size: int = Form(500, ge=50, le=10000),
    chunk_overlap: int = Form(50, ge=0, le=5000),
    tenant_id: str = Depends(get_tenant_id),
    db: AsyncSession = Depends(get_db),
):
    """Upload a text document, split into chunks, and store in the knowledge base."""
    import os

    # ── Validate chunk_overlap < chunk_size ──────────────────────
    if chunk_overlap >= chunk_size:
        raise HTTPException(
            400,
            f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size})",
        )

    # ── Validate file extension ──────────────────────────────────
    if not file.filename:
        raise HTTPException(400, "Filename is required")

    safe_name = _sanitize_filename(file.filename)
    ext = os.path.splitext(safe_name)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            400,
            f"Unsupported file type '{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # ── Verify source exists and is owned by the caller's tenant ─
    source_id = (source_id or "").strip() or None
    if source_id:
        result = await db.execute(
            select(KnowledgeSource).where(
                KnowledgeSource.id == source_id, KnowledgeSource.tenant_id == tenant_id
            )
        )
        source = result.scalar_one_or_none()
        if not source:
            raise HTTPException(404, f"Knowledge source '{source_id}' not found")
    else:
        base_name = (source_name or os.path.splitext(safe_name)[0]).strip()
        unique_name = await _unique_source_name(db, base_name, tenant_id)
        source = KnowledgeSource(
            name=unique_name,
            source_type="document",
            source_uri=safe_name,
            domain=domain,
            tenant_id=tenant_id,
            status="processing",
            metadata_={"filename": safe_name, "auto_created": True},
        )
        db.add(source)
        await db.flush()
        source_id = source.id

    # ── Read and parse file content ──────────────────────────────
    max_bytes = _max_upload_bytes()
    raw_bytes = await file.read(max_bytes + 1)
    if len(raw_bytes) > max_bytes:
        raise HTTPException(
            413,
            f"File is too large. Max knowledge upload size is {settings.knowledge_max_upload_mb} MB",
        )
    if not raw_bytes:
        raise HTTPException(400, "Uploaded file is empty")

    text = _extract_text_by_extension(raw_bytes, ext)

    if not text.strip():
        raise HTTPException(400, "Uploaded file is empty")

    # ── Split into chunks ────────────────────────────────────────
    chunks = _recursive_split(text, chunk_size, chunk_overlap)

    # ── Store chunks ─────────────────────────────────────────────
    created_chunks: list[KnowledgeChunk] = []
    for idx, chunk_content in enumerate(chunks):
        entity_key = _extract_entity_key(chunk_content)
        chunk = KnowledgeChunk(
            source_id=source_id,
            content=chunk_content,
            entity_key=entity_key,
            domain=domain,
            chunk_index=idx,
            metadata_={
                "filename": safe_name,
                "file_ext": ext,
                "file_size_bytes": len(raw_bytes),
                "chunk_size": chunk_size,
            },
        )
        db.add(chunk)
        created_chunks.append(chunk)

    # ── Update source chunk_count and status ─────────────────────
    source.chunk_count = (source.chunk_count or 0) + len(chunks)
    source.status = "ready"

    await db.flush()
    created_ids = [chunk.id for chunk in created_chunks if chunk.id]

    await db.commit()

    batch = [
        {"chunk_id": cid, "text": ct, "domain": domain}
        for cid, ct in zip(created_ids, chunks)
    ]
    background_tasks.add_task(_embed_batch_background, batch)

    return {
        "source_id": source_id,
        "filename": safe_name,
        "chunk_count": len(chunks),
        "chunk_ids": created_ids,
    }
